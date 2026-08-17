from __future__ import annotations

import html
import json
import re
import shutil
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SPEC_ROOT = ROOT.parent / "产品规格书"
SPEC_IMAGES = SPEC_ROOT / "_images"
ASSET_IMAGES = ROOT / "assets" / "images"
DOMAIN = "https://www.panpantechnology.com"


def supplier_file(*parts: str) -> Path:
    return ROOT.parent.joinpath(*parts)


CATEGORIES = [
    {
        "id": "indoor",
        "label": "Indoor Cleaning",
        "heading": "Compact indoor cleaning robots",
        "copy": "Autonomous sweep, vacuum, mop and scrub robots for offices, hotels, healthcare, retail, campuses and narrow indoor routes.",
        "products": ["acr-0350", "acr-0370", "acr-0440c", "acr-0440p", "acr-0520", "acr-0600", "acr-0670"],
    },
    {
        "id": "large-scrubbers",
        "label": "Large-Area Scrubbing",
        "heading": "Autonomous scrubbers for high-traffic floors",
        "copy": "Driverless large-area scrubbers for transport hubs, malls, factories, warehouses, exhibition venues and broad hard floors.",
        "products": ["acr-0800"],
    },
    {
        "id": "outdoor",
        "label": "Outdoor Sweeping",
        "heading": "Outdoor sweeping robots for campuses and municipal routes",
        "copy": "Wide-path autonomous sweepers for plazas, business parks, campuses, stations, scenic areas and municipal open spaces.",
        "products": ["acr-1200", "acr-1800"],
    },
    {
        "id": "amr",
        "label": "Warehouse AMR",
        "heading": "AMR platforms for industrial material handling",
        "copy": "Autonomous mobile robots for production feeding, warehouse transfer, heavy-payload delivery and multi-floor logistics.",
        "products": ["amr-0300", "amr-0600"],
    },
    {
        "id": "picking",
        "label": "Picking Robot",
        "heading": "Robotic picking and shelf-side handling",
        "copy": "Mobile manipulation robots for shelf picking, item transfer, tray loading and assisted fulfilment workflows.",
        "products": ["asr-0012"],
    },
    {
        "id": "facade",
        "label": "Facade Cleaning",
        "heading": "Glass facade cleaning robots",
        "copy": "Commercial facade robots for curtain-wall glass, high-rise exterior cleaning and safer building-maintenance workflows.",
        "products": ["facade-cleaning-robot"],
    },
]


BAD_MODEL_NAMES = {
    "ONE S55",
    "P060",
    "SC80",
    "PT90",
    "C2",
    "C2 PRO",
    "C3-Mini",
    "Q3-G",
    "Q3-W",
    "IQX70B",
    "XG",
    "YZ",
    "T300",
    "T600",
}


SUPPLEMENTAL_IMAGES = {
    "acr-0350": [
        (supplier_file("智绘科技 清洁机器人", "C3mini清洁机器人产品介绍2025-5.jpg"), "acr0350_extra1.jpg"),
        (supplier_file("智绘科技 清洁机器人", "C3mini清洁机器人产品介绍2025-10.jpg"), "acr0350_extra2.jpg"),
        (supplier_file("智绘科技 清洁机器人", "C3mini清洁机器人产品介绍2025-12.jpg"), "acr0350_extra3.jpg"),
    ],
    "acr-0370": [
        (supplier_file("浙江齐元机器人", "Q3-G 2D款 宣传册-中英-竖版-251106-v11.3.jpg"), "acr0370_extra1.jpg"),
        (supplier_file("浙江齐元机器人", "Q3-G 2D款 宣传册-中英-竖版-251106-v11.3-4.jpg"), "acr0370_extra2.jpg"),
        (supplier_file("浙江齐元机器人", "Q3-G 2D款 宣传册-中英-竖版-251106-v11.3-8.jpg"), "acr0370_extra3.jpg"),
    ],
    "acr-0440c": [
        (supplier_file("智绘科技 清洁机器人", "ALLYBOT-C2 EN-2024-5.jpg"), "acr0440c_extra1.jpg"),
        (supplier_file("智绘科技 清洁机器人", "ALLYBOT-C2 EN-2024-1.jpg"), "acr0440c_extra2.jpg"),
        (supplier_file("智绘科技 清洁机器人", "ALLYBOT-C2 EN-2024-10.jpg"), "acr0440c_extra3.jpg"),
        (supplier_file("智绘科技 清洁机器人", "ALLYBOT-C2 EN-2024-11.jpg"), "acr0440c_extra4.jpg"),
    ],
    "acr-0440p": [
        (supplier_file("智绘科技 清洁机器人", "C2pro机器人产品彩页0912-en-1.jpg"), "acr0440p_extra1.jpg"),
        (supplier_file("智绘科技 清洁机器人", "C2pro机器人产品彩页0912-en-2.jpg"), "acr0440p_extra2.jpg"),
        (supplier_file("智绘科技 清洁机器人", "C2pro机器人产品彩页0912-en-6.jpg"), "acr0440p_extra3.jpg"),
    ],
    "acr-0520": [
        (supplier_file("奇勃（深圳）科技", "原始照片", "微信图片_20250904115004.png"), "acr0520_extra1.png"),
        (supplier_file("奇勃（深圳）科技", "原始照片", "微信图片_20250904114949.jpg"), "acr0520_extra2.jpg"),
        (supplier_file("奇勃（深圳）科技", "原始照片", "微信图片_20250904114959.jpg"), "acr0520_extra3.jpg"),
        (supplier_file("奇勃（深圳）科技", "原始照片", "微信图片_20250904114900.jpg"), "acr0520_extra4.jpg"),
    ],
    "acr-0600": [
        (supplier_file("神州云海", "主图1.jpg"), "acr0600_extra1.jpg"),
        (supplier_file("神州云海", "主图10.jpg"), "acr0600_extra2.jpg"),
        (supplier_file("神州云海", "神州云海清洁机器人介绍 IQX70B-6.jpg"), "acr0600_extra3.jpg"),
        (supplier_file("神州云海", "主图13.jpg"), "acr0600_extra4.jpg"),
    ],
    "acr-0670": [
        (supplier_file("浙江齐元机器人", "Q3-G+W宣传册-中英-竖版-250925-v8.5.jpg"), "acr0670_extra1.jpg"),
        (supplier_file("浙江齐元机器人", "Q3-G+W宣传册-中英-竖版-250925-v8.5-2.jpg"), "acr0670_extra2.jpg"),
        (supplier_file("浙江齐元机器人", "Q3-G+W宣传册-中英-竖版-250925-v8.5-3.jpg"), "acr0670_extra3.jpg"),
    ],
    "acr-1200": [
        (supplier_file("九天创新 毛世鑫", "XG 1st Product Brochure-1.jpg"), "acr1200_extra1.jpg"),
        (supplier_file("九天创新 毛世鑫", "XG 1st Product Brochure-2.jpg"), "acr1200_extra2.jpg"),
        (supplier_file("九天创新 毛世鑫", "XG 1st Product Brochure-4.jpg"), "acr1200_extra3.jpg"),
        (supplier_file("九天创新 毛世鑫", "璇光智慧清洁机器人（中文--电子版2.0）-2.jpg"), "acr1200_extra4.jpg"),
    ],
    "acr-1800": [
        (supplier_file("九天创新 毛世鑫", "YZ Product Brochure-1.jpg"), "acr1800_extra1.jpg"),
        (supplier_file("九天创新 毛世鑫", "YZ Product Brochure-2.jpg"), "acr1800_extra2.jpg"),
        (supplier_file("九天创新 毛世鑫", "YZ Product Brochure-3.jpg"), "acr1800_extra3.jpg"),
        (supplier_file("九天创新 毛世鑫", "YZ Product Brochure-4.jpg"), "acr1800_extra4.jpg"),
        (supplier_file("九天创新 毛世鑫", "YZ Product Brochure-5.jpg"), "acr1800_extra5.jpg"),
    ],
    "amr-0300": [
        (supplier_file("深圳市普渡科技", "T300-20250917T093742Z-1-001", "T300", "picture", "工厂场景(1).jpg"), "amr0300_extra1.jpg"),
        (supplier_file("深圳市普渡科技", "T300-20250917T093742Z-1-001", "T300", "Accessories", "双层辊筒高效果2正侧视 拷贝.png"), "amr0300_extra2.png"),
        (supplier_file("深圳市普渡科技", "T300-20250917T093742Z-1-001", "T300", "Accessories", "单层辊筒最高效果2 拷贝.png"), "amr0300_extra3.png"),
    ],
    "amr-0600": [
        (supplier_file("深圳市普渡科技", "T600 promotional materials package-20250917T093805Z-1-001", "T600 promotional materials package", "T600-pictures产品图示 (3)", "柔性部署.jpg"), "amr0600_extra1.jpg"),
        (supplier_file("深圳市普渡科技", "T600 promotional materials package-20250917T093805Z-1-001", "T600 promotional materials package", "T600-pictures产品图示 (3)", "主图场景.jpg"), "amr0600_extra2.jpg"),
        (supplier_file("深圳市普渡科技", "T600 promotional materials package-20250917T093805Z-1-001", "T600 promotional materials package", "T600-pictures产品图示 (3)", "切换-立柱.jpg"), "amr0600_extra3.jpg"),
        (supplier_file("深圳市普渡科技", "T600 promotional materials package-20250917T093805Z-1-001", "T600 promotional materials package", "T600-pictures产品图示 (3)", "切换-潜伏.jpg"), "amr0600_extra4.jpg"),
    ],
}


EXCLUDED_GALLERY_IMAGES = {
    "acr0350_extra1.jpg",
    "acr0350_extra2.jpg",
    "acr0440c_extra1.jpg",
    "acr0440c_extra3.jpg",
    "acr0440c_extra4.jpg",
    "acr0440p_extra1.jpg",
    "acr0440p_extra2.jpg",
    "acr0440p_extra3.jpg",
    "acr0600_extra1.jpg",
    "acr0600_extra3.jpg",
    "acr0670_v3.jpg",
    "acr0670_extra1.jpg",
    "acr0670_extra2.jpg",
    "acr0670_extra3.jpg",
    "acr1200_extra3.jpg",
    "acr1200_extra4.jpg",
    "acr1800_extra2.jpg",
}


MANUAL_PRODUCTS = [
    {
        "slug": "asr-0012",
        "model": "ASR-0012",
        "name": "Shelf Picking and Mobile Manipulation Robot",
        "subtitle": "Picking Robot",
        "description": "The ASR-0012 is a mobile manipulation robot for shelf-side picking, tote handling and light item transfer in retail, warehouse and hospitality workflows. It combines an autonomous mobile base with a vertical lift column, perception sensors and a multi-axis robotic arm, allowing operators to trial item retrieval, shelf replenishment and delivery handoff without rebuilding the site around fixed automation.",
        "metrics": [
            {"value": "Mobile", "label": "autonomous base"},
            {"value": "Arm", "label": "robotic picking"},
            {"value": "Shelf", "label": "retail and warehouse"},
            {"value": "Pilot", "label": "workflow validation"},
        ],
        "advantages": [
            {
                "title": "Mobile picking at the shelf edge",
                "copy": "The robot brings a manipulator to the shelf, bin or table rather than requiring every item to move through a fixed picking cell.",
            },
            {
                "title": "Suitable for pilots and workflow proofing",
                "copy": "Use it to test tote presentation, shelf approach, item handoff, operator supervision and exception handling before committing to a larger automation program.",
            },
            {
                "title": "Autonomous base with vertical reach",
                "copy": "The mobile base and lift column support access to different shelf heights, while the arm handles light item retrieval and placement tasks.",
            },
            {
                "title": "Shelf, tray and service scenarios",
                "copy": "Typical trials include convenience retail shelves, storage and picking zones, restaurant table service, tray loading and material handoff points.",
            },
            {
                "title": "Computer vision assisted operation",
                "copy": "Camera and perception modules help the robot locate the work area and align with products, totes or fixtures during supervised deployment.",
            },
            {
                "title": "Works alongside AMR transport",
                "copy": "For larger sites, combine picking or item handling with AMR transport pages to build a complete fulfilment path from shelf to workstation.",
            },
        ],
        "specs": [
            {"key": "Robot type", "value": "Mobile manipulation / shelf picking robot"},
            {"key": "Model", "value": "ASR-0012"},
            {"key": "Main modules", "value": "Autonomous mobile base, lift column, robotic arm, vision and perception sensors"},
            {"key": "Primary tasks", "value": "Shelf picking, tote handoff, tray loading, item transfer and assisted service workflows"},
            {"key": "Deployment mode", "value": "Pilot or project configuration after site and SKU review"},
            {"key": "Best-fit objects", "value": "Light retail goods, packaged items, trays and repeatable handoff tasks"},
            {"key": "Workflow inputs", "value": "Shelf dimensions, item size and weight, bin/tote style, route width and handoff height"},
            {"key": "Navigation", "value": "Autonomous mobile navigation with supervised workflow setup"},
            {"key": "Perception", "value": "Vision-assisted work-area recognition and object approach"},
            {"key": "Integration", "value": "Can be evaluated with AMR delivery, shelf replenishment, goods-to-person and line-side workflows"},
            {"key": "Final configuration", "value": "Confirm through PanPanTech contact form with site photos and task video"},
        ],
        "applications": [
            "Retail shelf picking",
            "Warehouse picking zones",
            "Tote and tray handling",
            "Restaurant service support",
            "Shelf replenishment trials",
            "R&D pilot projects",
        ],
        "images": [
            "/assets/images/asr0012_v1.jpg",
            "/assets/images/asr0012_v2.jpg",
            "/assets/images/asr0012_v3.jpg",
        ],
        "category": "Picking Robot",
        "source": "PanPanTech ASR-0012 image pack",
    },
    {
        "slug": "facade-cleaning-robot",
        "model": "Facade Cleaning Robot",
        "name": "for Glass Curtain Walls",
        "subtitle": "Glass Facade Cleaning",
        "description": "The PanPanTech facade cleaning robot page covers commercial glass curtain-wall cleaning for office towers, malls, hotels, campuses and property-service teams. The system is positioned as a project solution rather than a stand-alone consumer window device: building height, glass geometry, mullion layout, access points, wind exposure, water workflow and local work-at-height rules must be reviewed before model confirmation.",
        "metrics": [
            {"value": "Glass", "label": "curtain walls"},
            {"value": "High-rise", "label": "exterior cleaning"},
            {"value": "Hybrid", "label": "robot plus operator"},
            {"value": "Pilot", "label": "site assessment"},
        ],
        "advantages": [
            {
                "title": "Designed around facade workflow",
                "copy": "Commercial facade cleaning depends on access, anchoring, surface geometry, water management and recovery procedures, so the page guides customers to share building information first.",
            },
            {
                "title": "Reduces repetitive work-at-height exposure",
                "copy": "On suitable repeated glass zones, the robot can move routine cleaning away from rope-only manual work and into a more controlled machine-assisted workflow.",
            },
            {
                "title": "Works with hybrid cleaning plans",
                "copy": "Robots handle regular glass areas while trained technicians remain responsible for setup, inspection, corners, signage, recesses and other difficult facade details.",
            },
            {
                "title": "Supports multiple facade materials",
                "copy": "Project assessment can cover pure flat glass, framed curtain walls, stone curtain walls, metal curtain walls and composite facade zones.",
            },
            {
                "title": "Pilot before rollout",
                "copy": "A representative pilot zone verifies cleaning quality, setup time, water supply, operating limits, recovery path and manual follow-up requirements.",
            },
            {
                "title": "Procurement-ready information request",
                "copy": "The page asks customers to submit facade drawings, photos, cleaning frequency, working hours and safety requirements through the website form.",
            },
        ],
        "specs": [
            {"key": "Robot type", "value": "Commercial glass facade / curtain-wall cleaning robot"},
            {"key": "Use case", "value": "Exterior building glass and repeated facade cleaning workflows"},
            {"key": "Typical projects", "value": "Office towers, commercial complexes, hotels, campuses, hospitals and industrial parks"},
            {"key": "Suitable surfaces", "value": "Flat glass, framed glass curtain walls, stone curtain walls, metal curtain walls and composite facades after assessment"},
            {"key": "Deployment basis", "value": "Building-specific assessment, pilot cleaning and safety review"},
            {"key": "Required inputs", "value": "Facade photos, drawings, building height, access points, mullion spacing, water/power plan and local safety rules"},
            {"key": "Operating plan", "value": "Robot-assisted cleaning with trained operator supervision and manual support for special details"},
            {"key": "Safety planning", "value": "Work-at-height procedure, exclusion zone, weather limit, inspection checklist and emergency recovery plan"},
            {"key": "Final configuration", "value": "Confirm through PanPanTech contact form before final deployment planning"},
        ],
        "applications": [
            "Glass curtain walls",
            "Office towers",
            "Hotels and malls",
            "Commercial complexes",
            "Campus buildings",
            "Property service teams",
        ],
        "images": [
            "/assets/images/facade-window-cleaning-robot-card.jpg",
            "/assets/images/facade-cleaning-robot-new.jpg",
            "/assets/images/facade-robot-building-assessment.jpg",
            "/assets/images/facade-robot-commercial-glass-wall.jpg",
        ],
        "category": "Facade Cleaning",
        "source": "Facade cleaning supplier materials",
    },
]


def esc(value: object) -> str:
    return html.escape(str(value), quote=True)


def clean_public_copy(value: str) -> str:
    replacements = {
        "fewer machines to buy and maintain": "a simpler equipment set to operate and maintain",
        "directly reducing labour and supervision cost": "reducing manual sweeping and supervision workload",
        "reducing labour and supervision costs": "reducing manual sweeping and supervision workload",
        "cost": "operating workload",
        "costs": "operating workload",
        "price": "configuration",
        "pricing": "configuration",
        "offers": "provides",
        "offer": "provide",
    }
    cleaned = value
    for source, target in replacements.items():
        cleaned = re.sub(re.escape(source), target, cleaned, flags=re.I)
    return cleaned


def slug_from_file(path: Path) -> str:
    match = re.search(r"PanPanTech_([A-Z]+)-(\d+[A-Z]?)_", path.name)
    if not match:
        raise ValueError(f"Cannot derive slug from {path.name}")
    return f"{match.group(1).lower()}-{match.group(2).lower()}"


def image_stem(slug: str) -> str:
    return slug.replace("-", "")


def split_title(title: str) -> tuple[str, str]:
    parts = re.split(r"\s+", title.replace("\u3000", " "), maxsplit=1)
    model = parts[0].strip()
    name = parts[1].strip() if len(parts) > 1 else model
    return model, name


def table_text(table, row: int, col: int) -> str:
    return table.cell(row, col).text.strip()


def extract_products() -> list[dict]:
    products = []
    for path in sorted(SPEC_ROOT.glob("*Product_Specification_EN_v1.0.docx")):
        slug = slug_from_file(path)
        document = Document(path)
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        title = next((p for p in paragraphs if p.startswith(("ACR-", "AMR-"))), "")
        if not title:
            continue
        title_index = paragraphs.index(title)
        subtitle = paragraphs[title_index + 1] if title_index + 1 < len(paragraphs) else ""
        model, name = split_title(title)
        tables = document.tables
        description = clean_public_copy(table_text(tables[0], 0, 0))

        metric_values = [cell.text.strip() for cell in tables[1].rows[0].cells]
        metric_labels = [cell.text.strip() for cell in tables[1].rows[1].cells]
        metrics = [{"value": value, "label": label} for value, label in zip(metric_values, metric_labels)]

        advantages = []
        for row in tables[2].rows:
            title_cell = row.cells[0].text.strip()
            copy_cell = row.cells[1].text.strip()
            if title_cell and copy_cell:
                advantages.append({"title": clean_public_copy(title_cell), "copy": clean_public_copy(copy_cell)})

        specs = []
        seen_section_headers = set()
        for row in tables[4].rows:
            key = row.cells[0].text.strip()
            value = row.cells[1].text.strip()
            if not key or not value:
                continue
            if key == value:
                seen_section_headers.add(key)
                continue
            specs.append({"key": clean_public_copy(key), "value": clean_public_copy(value)})

        applications = []
        for row in tables[5].rows:
            for cell in row.cells:
                value = cell.text.strip()
                if value:
                    applications.append(clean_public_copy(value))

        stem = image_stem(slug)
        images = sorted(SPEC_IMAGES.glob(f"{stem}_v*.jpg"))
        for image in images:
            if image.name in EXCLUDED_GALLERY_IMAGES:
                continue
            shutil.copy2(image, ASSET_IMAGES / image.name)
        image_urls = [f"/assets/images/{image.name}" for image in images if image.name not in EXCLUDED_GALLERY_IMAGES]
        for source, output_name in SUPPLEMENTAL_IMAGES.get(slug, []):
            if output_name in EXCLUDED_GALLERY_IMAGES:
                continue
            if not source.exists():
                raise FileNotFoundError(f"Supplemental image missing for {slug}: {source}")
            shutil.copy2(source, ASSET_IMAGES / output_name)
            image_urls.append(f"/assets/images/{output_name}")

        category = next(category["label"] for category in CATEGORIES if slug in category["products"])
        products.append(
            {
                "slug": slug,
                "model": model,
                "name": name,
                "subtitle": subtitle,
                "description": description,
                "metrics": metrics,
                "advantages": advantages,
                "specs": specs,
                "applications": applications,
                "images": image_urls,
                "category": category,
                "source": path.name,
            }
        )
    for product in MANUAL_PRODUCTS:
        for src in product["images"]:
            source = SPEC_IMAGES / Path(src).name
            target = ROOT / src.lstrip("/")
            if source.exists() and not target.exists():
                shutil.copy2(source, target)
        products.append(product)
    return products


def org_json_ld() -> str:
    return json.dumps(
        {
            "@context": "https://schema.org",
            "@type": "Organization",
            "name": "PanPanTech",
            "url": DOMAIN,
            "logo": f"{DOMAIN}/assets/images/panpantech-logo.png",
            "email": "info@panpantechnology.com",
            "telephone": "+86-13925118851",
            "image": f"{DOMAIN}/assets/images/panpantech-social-card.jpg",
        },
        ensure_ascii=False,
        separators=(",", ":"),
    )


def breadcrumb_json_ld(items: list[tuple[str, str]]) -> str:
    return json.dumps(
        {
            "@context": "https://schema.org",
            "@type": "BreadcrumbList",
            "itemListElement": [
                {"@type": "ListItem", "position": index + 1, "name": name, "item": f"{DOMAIN}{url}"}
                for index, (name, url) in enumerate(items)
            ],
        },
        ensure_ascii=False,
        separators=(",", ":"),
    )


def product_json_ld(product: dict) -> str:
    return json.dumps(
        {
            "@context": "https://schema.org",
            "@type": "Product",
            "name": f"{product['model']} {product['name']}",
            "image": f"{DOMAIN}{product['images'][0]}" if product["images"] else f"{DOMAIN}/assets/images/panpantech-social-card.jpg",
            "description": product["description"],
            "brand": {"@type": "Brand", "name": "PanPanTech"},
            "sku": product["model"],
            "category": product["category"],
            "additionalProperty": [
                {"@type": "PropertyValue", "name": item["key"], "value": item["value"]}
                for item in product["specs"][:12]
            ],
        },
        ensure_ascii=False,
        separators=(",", ":"),
    )


def head(title: str, description: str, canonical: str, extra_json_ld: list[str]) -> str:
    json_ld = "\n".join(f'<script type="application/ld+json">{block}</script>' for block in [org_json_ld(), *extra_json_ld])
    return f"""<!doctype html>
<html><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<link rel="preconnect" href="https://fonts.googleapis.com"><link rel="preconnect" href="https://fonts.gstatic.com" crossorigin="anonymous"><link href="https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@400;500;600;700&amp;family=Instrument+Sans:ital,wght@0,400;0,500;0,600;1,400&amp;family=IBM+Plex+Mono:wght@400;500;600&amp;display=swap" rel="stylesheet" media="print" onload="this.media='all'"><noscript><link href="https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@400;500;600;700&amp;family=Instrument+Sans:ital,wght@0,400;0,500;0,600;1,400&amp;family=IBM+Plex+Mono:wght@400;500;600&amp;display=swap" rel="stylesheet"></noscript>
<title>{esc(title)}</title><meta name="description" content="{esc(description)}"><link rel="canonical" href="{DOMAIN}{canonical}">
<meta property="og:type" content="website"><meta property="og:title" content="{esc(title)}"><meta property="og:description" content="{esc(description)}"><meta property="og:url" content="{DOMAIN}{canonical}"><meta property="og:image" content="{DOMAIN}/assets/images/panpantech-social-card.jpg">
<meta name="twitter:card" content="summary_large_image"><meta name="twitter:image" content="{DOMAIN}/assets/images/panpantech-social-card.jpg">
{json_ld}
<style>
:root {{ --ink:#0b0f1a; --muted:#6a7386; --line:#e2e6ed; --blue:#0e5fd9; --bg:#f5f6f8; --panel:#ffffff; --dark:#070a12; }}
html, body {{ margin:0; padding:0; background:var(--bg); }}
* {{ box-sizing:border-box; }}
body {{ font-family:"Instrument Sans", system-ui, sans-serif; color:var(--muted); background:var(--bg); }}
a {{ color:inherit; }}
.container {{ max-width:1360px; margin:0 auto; padding-left:40px; padding-right:40px; }}
.mono {{ font-family:"IBM Plex Mono", monospace; }}
.brand {{ font-family:"Space Grotesk", sans-serif; font-size:21px; font-weight:700; letter-spacing:0; color:var(--ink); text-decoration:none; }}
.brand span {{ color:var(--blue); }}
.nav {{ position:sticky; top:0; z-index:90; background:rgba(245,246,248,.9); backdrop-filter:blur(14px); border-bottom:1px solid var(--line); }}
.nav-inner {{ height:72px; display:flex; align-items:center; justify-content:space-between; gap:28px; }}
.nav-links {{ display:flex; align-items:center; gap:34px; }}
.nav-links a {{ font-size:14px; font-weight:500; color:#3d4453; text-decoration:none; }}
.nav-links .active {{ font-weight:600; color:var(--ink); border-bottom:2px solid var(--blue); padding-bottom:2px; }}
.pill-btn {{ display:inline-flex; align-items:center; justify-content:center; gap:10px; border-radius:999px; text-decoration:none; font-weight:600; white-space:nowrap; }}
.pill-dark {{ background:var(--ink); color:#fff; padding:12px 22px; font-size:13.5px; }}
.pill-blue {{ background:var(--blue); color:#fff; padding:15px 28px; font-size:14.5px; }}
.pill-outline {{ border:1px solid var(--ink); color:var(--ink); padding:15px 28px; font-size:14.5px; }}
.hero {{ background:var(--dark); color:#fff; }}
.hero-inner {{ padding-top:84px; padding-bottom:72px; }}
.eyebrow {{ margin:0 0 18px; font-family:"IBM Plex Mono", monospace; font-size:12px; letter-spacing:.22em; text-transform:uppercase; color:var(--blue); }}
h1, h2, h3 {{ font-family:"Space Grotesk", sans-serif; color:var(--ink); letter-spacing:0; }}
.hero h1 {{ margin:0; max-width:920px; color:#fff; font-size:58px; line-height:1.04; font-weight:600; }}
.hero-copy {{ margin:24px 0 0; max-width:620px; font-size:17px; line-height:1.6; color:#b9c4d8; }}
.category-tabs {{ display:flex; flex-wrap:wrap; gap:10px; margin-top:40px; }}
.category-tabs a {{ display:inline-flex; align-items:center; gap:8px; border:1px solid rgba(255,255,255,.24); color:#fff; font-size:13.5px; font-weight:500; padding:10px 18px; border-radius:999px; text-decoration:none; }}
.band {{ border-bottom:1px solid var(--line); }}
.band.alt {{ background:#edeff3; }}
.band-inner {{ padding-top:84px; padding-bottom:94px; }}
.section-head {{ display:flex; align-items:flex-end; justify-content:space-between; gap:40px; margin-bottom:42px; }}
.section-head h2 {{ margin:0; font-size:40px; line-height:1.08; font-weight:600; max-width:660px; }}
.section-head p:last-child {{ margin:0 0 4px; max-width:430px; font-size:15px; line-height:1.6; }}
.grid {{ display:grid; grid-template-columns:repeat(3, minmax(0, 1fr)); gap:22px; }}
.card {{ display:flex; flex-direction:column; background:#fff; border:1px solid var(--line); border-radius:14px; overflow:hidden; text-decoration:none; color:inherit; transition:transform 180ms, box-shadow 180ms; }}
.card:hover {{ transform:translateY(-3px); box-shadow:0 18px 45px rgba(20,30,50,.08); }}
.media {{ aspect-ratio:4/3; background:#f0f2f5; overflow:hidden; }}
.media img {{ width:100%; height:100%; object-fit:contain; display:block; padding:12px; }}
.card-body {{ display:flex; flex-direction:column; flex:1; padding:25px 27px 27px; }}
.tag {{ font-family:"IBM Plex Mono", monospace; font-size:11px; letter-spacing:.14em; text-transform:uppercase; color:var(--blue); }}
.model {{ margin-top:10px; font-family:"Space Grotesk", sans-serif; font-size:22px; font-weight:600; color:var(--ink); }}
.card-copy {{ margin-top:8px; font-size:14px; line-height:1.55; color:var(--muted); }}
.metric-row {{ display:flex; justify-content:space-between; gap:16px; border-top:1px solid var(--line); padding-top:16px; margin-top:auto; }}
.metric-row span {{ font-family:"IBM Plex Mono", monospace; font-size:12px; color:var(--muted); }}
.metric-row strong {{ font-family:"IBM Plex Mono", monospace; font-size:12px; color:var(--ink); }}
.crumbs {{ display:flex; gap:10px; padding-top:20px; font-family:"IBM Plex Mono", monospace; font-size:12px; color:#9aa3b5; }}
.crumbs a {{ text-decoration:none; color:#9aa3b5; }}
.product-hero {{ border-bottom:1px solid var(--line); }}
.product-hero-inner {{ padding-top:48px; padding-bottom:88px; display:grid; grid-template-columns:1.05fr 1fr; gap:64px; align-items:center; }}
.hero-image {{ border-radius:14px; overflow:hidden; border:1px solid var(--line); aspect-ratio:1/1; background:#fff; }}
.hero-image img {{ width:100%; height:100%; object-fit:contain; display:block; padding:18px; }}
.product-title {{ margin:0; font-size:52px; line-height:1.04; font-weight:600; }}
.lead {{ margin:22px 0 0; font-size:17px; line-height:1.65; color:var(--muted); }}
.stats {{ display:grid; grid-template-columns:repeat(4, 1fr); border-top:1px solid var(--line); margin-top:34px; padding-top:24px; gap:0; }}
.stat {{ border-right:1px solid var(--line); padding:0 18px; min-width:0; }}
.stat:first-child {{ padding-left:0; }}
.stat:last-child {{ border-right:0; }}
.stat strong {{ display:block; font-family:"Space Grotesk", sans-serif; font-size:28px; line-height:1.1; color:var(--ink); }}
.stat span {{ display:block; margin-top:6px; font-family:"IBM Plex Mono", monospace; font-size:10.5px; letter-spacing:.12em; text-transform:uppercase; color:#9aa3b5; }}
.cta-row {{ display:flex; gap:12px; flex-wrap:wrap; margin-top:34px; }}
.two-col {{ display:grid; grid-template-columns:1.35fr 1fr; gap:64px; }}
.advantage-list {{ margin-top:34px; display:flex; flex-direction:column; border-top:1px solid #d8dde6; }}
.advantage {{ display:grid; grid-template-columns:56px 1fr; gap:20px; padding:22px 0; border-bottom:1px solid #d8dde6; }}
.advantage strong {{ display:block; font-family:"Space Grotesk", sans-serif; font-size:18px; color:var(--ink); }}
.advantage p {{ margin:6px 0 0; font-size:14.5px; line-height:1.6; }}
.chips {{ display:flex; flex-wrap:wrap; gap:8px; margin-top:22px; border-top:1px solid var(--line); padding-top:20px; }}
.chips span {{ font-family:"IBM Plex Mono", monospace; font-size:12px; color:#3d4453; border:1px solid #d8dde6; padding:7px 13px; border-radius:999px; }}
.panel {{ background:#fff; border:1px solid var(--line); border-radius:14px; padding:34px 36px; align-self:start; }}
.spec-grid {{ display:grid; grid-template-columns:1fr 1fr; gap:0 58px; border-top:1px solid #d8dde6; }}
.spec {{ display:grid; grid-template-columns:210px 1fr; gap:24px; padding:17px 0; border-bottom:1px solid var(--line); align-items:baseline; }}
.spec span:first-child {{ font-family:"IBM Plex Mono", monospace; font-size:12px; letter-spacing:.08em; text-transform:uppercase; color:#9aa3b5; }}
.spec span:last-child {{ font-size:15px; font-weight:500; color:var(--ink); }}
.gallery {{ display:grid; grid-template-columns:repeat(3, 1fr); gap:22px; margin-top:42px; }}
.gallery figure {{ margin:0; background:#fff; border:1px solid var(--line); border-radius:14px; overflow:hidden; }}
.gallery img {{ width:100%; aspect-ratio:4/3; object-fit:contain; display:block; padding:14px; }}
.gallery-intro {{ display:flex; align-items:flex-end; justify-content:space-between; gap:40px; margin-bottom:34px; }}
.gallery-intro h2 {{ margin:0; font-size:38px; line-height:1.1; }}
.gallery-intro p {{ margin:0; max-width:430px; font-size:15px; line-height:1.6; }}
.dark-cta {{ background:var(--dark); color:#fff; }}
.dark-cta .container {{ padding-top:90px; padding-bottom:90px; display:flex; align-items:center; justify-content:space-between; gap:60px; }}
.dark-cta h2 {{ margin:0; color:#fff; max-width:740px; font-size:44px; line-height:1.08; font-weight:600; }}
.dark-cta p {{ margin:18px 0 0; max-width:560px; font-size:16px; line-height:1.6; color:#8fa3c8; }}
footer {{ background:var(--dark); border-top:1px solid rgba(255,255,255,.1); color:#b9c4d8; }}
.footer-inner {{ padding-top:60px; padding-bottom:40px; display:grid; grid-template-columns:1.4fr 1fr 1fr 1.25fr; gap:42px; }}
footer a {{ color:#b9c4d8; text-decoration:none; }}
footer h3 {{ margin:0 0 12px; font-family:"IBM Plex Mono", monospace; font-size:11px; letter-spacing:.16em; text-transform:uppercase; color:#5c6b85; }}
footer p, footer a {{ font-size:14px; line-height:1.65; }}
.copyright {{ border-top:1px solid rgba(255,255,255,.1); padding:22px 0 28px; font-family:"IBM Plex Mono", monospace; font-size:12px; color:#5c6b85; }}
@media (max-width: 900px) {{
  .container {{ padding-left:20px; padding-right:20px; }}
  .nav-inner {{ height:auto; padding-top:18px; padding-bottom:18px; align-items:flex-start; }}
  .nav-links {{ display:none; }}
  .hero h1 {{ font-size:42px; }}
  .section-head, .dark-cta .container {{ display:block; }}
  .grid, .product-hero-inner, .two-col, .spec-grid, .gallery, .footer-inner {{ grid-template-columns:1fr; }}
  .stats {{ grid-template-columns:1fr 1fr; gap:18px 0; }}
  .stat:nth-child(2) {{ border-right:0; }}
  .product-title {{ font-size:40px; }}
  .spec {{ grid-template-columns:1fr; gap:6px; }}
}}
</style><link rel="icon" href="/favicon.ico" sizes="any"><link rel="icon" href="/assets/images/favicon-48x48.png" type="image/png" sizes="48x48"><link rel="apple-touch-icon" href="/assets/images/apple-touch-icon.png"><link rel="stylesheet" href="/assets/css/responsive.css"><meta name="theme-color" content="#0E5FD9"></head>"""


def header() -> str:
    return """<body><header class="nav"><div class="container nav-inner">
<a href="/" class="brand">PanPan<span>Tech</span></a>
<nav class="nav-links">
<a href="/solutions/">Solutions</a><a href="/products/" class="active">Smart Robots</a><a href="/technology/">Technology</a><a href="/manufacturing/">Manufacturing</a><a href="/partners/">Partners</a><a href="/blog/">Blog</a><a href="/about/">About</a>
</nav>
<a href="/request-a-quote/" class="pill-btn pill-dark">Contact Us<span class="mono">→</span></a>
</div></header>"""


def footer() -> str:
    return """<footer><div class="container footer-inner">
<div><div class="brand" style="color:#fff">PanPan<span>Tech</span></div><p>Retail AIoT, smart robotics, and manufacturing platform for global B2B buyers, integrators, distributors, and OEM partners.</p></div>
<div><h3>Robots</h3><p><a href="/products/#indoor">Indoor cleaning</a><br><a href="/products/#large-scrubbers">Large-area scrubbing</a><br><a href="/products/#outdoor">Outdoor sweeping</a><br><a href="/products/#amr">Warehouse AMR</a><br><a href="/products/#picking">Picking robot</a><br><a href="/products/#facade">Facade cleaning</a></p></div>
<div><h3>Company</h3><p><a href="/technology/">Technology</a><br><a href="/manufacturing/">Manufacturing</a><br><a href="/about/">About</a><br><a href="/request-a-quote/">Contact form</a></p></div>
<div><h3>Contact</h3><p>Building A1, Yuexiu iPARK Yuegang Zhigu, Nansha District, Guangzhou, Guangdong Province, China</p><p><a href="mailto:info@panpantechnology.com">info@panpantechnology.com</a><br><a href="tel:+8613925118851">+86-13925118851</a></p></div>
</div><div class="container copyright">&copy; 2026 PanPanTech&reg; | Retail AIoT &middot; Smart Robotics &middot; Manufacturing | Guangzhou PanPanTech Co., Ltd. All Rights Reserved.</div></footer><script src="/assets/js/site.js" defer></script></body></html>"""


def card(product: dict) -> str:
    metrics = product["metrics"]
    return f"""<a href="/products/{product['slug']}/" class="card">
<div class="media"><img src="{esc(product['images'][0])}" alt="{esc(product['model'])} {esc(product['name'])}" loading="lazy"></div>
<div class="card-body"><span class="tag">{esc(product['category'])}</span><span class="model">{esc(product['model'])}</span>
<span class="card-copy">{esc(product['name'])}. {esc(product['description'].split('.')[0])}.</span>
<div class="metric-row"><span>{esc(metrics[0]['value'])} {esc(metrics[0]['label'])}</span><strong>{esc(metrics[1]['value'])}</strong></div></div></a>"""


def products_index(products: list[dict]) -> str:
    by_slug = {product["slug"]: product for product in products}
    title = "PanPanTech Smart Robots | Cleaning Robots and Warehouse AMRs"
    description = "Explore PanPanTech ACR cleaning robots and AMR transport robots by task, payload, cleaning capacity, runtime, navigation and application fit."
    body = [
        head(title, description, "/products/", [breadcrumb_json_ld([("Home", "/"), ("Products", "/products/")])]),
        header(),
        f"""<section class="hero"><div class="container hero-inner">
<p class="eyebrow">Product Range</p><h1>PanPanTech ACR cleaning robots and AMR transport platforms</h1>
<p class="hero-copy">Compare verified PanPanTech model names and specifications from the latest product datasheets. For configuration details, contact the PanPanTech team through the website form.</p>
<div class="category-tabs">{''.join(f'<a href="#{c["id"]}"><span class="mono">{len(c["products"]):02d}</span>{esc(c["label"])}</a>' for c in CATEGORIES)}</div>
</div></section>
<span id="scrubbing" style="position:absolute"></span><span id="facade" style="position:absolute"></span><span id="service" style="position:absolute"></span>""",
    ]
    for index, category in enumerate(CATEGORIES):
        body.append(
            f"""<section id="{category['id']}" class="band {'alt' if index % 2 else ''}"><div class="container band-inner">
<div class="section-head"><div><p class="eyebrow">{index + 1:02d} / {esc(category['label'])}</p><h2>{esc(category['heading'])}</h2></div><p>{esc(category['copy'])}</p></div>
<div class="grid">{''.join(card(by_slug[slug]) for slug in category['products'])}</div>
</div></section>"""
        )
    body.append(
        """<section class="dark-cta"><div class="container"><div><h2>Need a model recommendation for your facility?</h2><p>Send floor area, surface type, working hours, expected route and destination country. We will match the right PanPanTech model and configuration.</p></div><a href="/request-a-quote/" class="pill-btn pill-blue">Contact the team<span class="mono">→</span></a></div></section>"""
    )
    body.append(footer())
    return "\n".join(body)


def product_page(product: dict) -> str:
    title = f"{product['model']} {product['name']} | PanPanTech"
    description = product["description"][:215]
    canonical = f"/products/{product['slug']}/"
    metrics = "".join(
        f"""<div class="stat"><strong>{esc(metric['value'])}</strong><span>{esc(metric['label'])}</span></div>"""
        for metric in product["metrics"]
    )
    advantages = "".join(
        f"""<div class="advantage"><span class="mono">/{index:02d}</span><div><strong>{esc(item['title'])}</strong><p>{esc(item['copy'])}</p></div></div>"""
        for index, item in enumerate(product["advantages"][:7], 1)
    )
    specs = "".join(
        f"""<div class="spec"><span>{esc(item['key'])}</span><span>{esc(item['value'])}</span></div>"""
        for item in product["specs"][:28]
    )
    applications = "".join(f"<span>{esc(app)}</span>" for app in product["applications"])
    gallery = "".join(
        f"""<figure><img src="{esc(src)}" alt="{esc(product['model'])} product view {index}" loading="lazy"></figure>"""
        for index, src in enumerate(product["images"], 1)
    )
    body = [
        head(
            title,
            description,
            canonical,
            [
                breadcrumb_json_ld([("Home", "/"), ("Products", "/products/"), (product["model"], canonical)]),
                product_json_ld(product),
            ],
        ),
        header(),
        f"""<div class="container crumbs"><a href="/">Home</a><span>/</span><a href="/products/">Products</a><span>/</span><span>{esc(product['model'])}</span></div>
<section class="product-hero"><div class="container product-hero-inner">
<div class="hero-image"><img src="{esc(product['images'][0])}" alt="{esc(product['model'])} {esc(product['name'])}" loading="eager"></div>
<div><p class="eyebrow">{esc(product['category'])}</p><h1 class="product-title">{esc(product['model'])} {esc(product['name'])}</h1><p class="lead">{esc(product['description'])}</p>
<div class="stats">{metrics}</div><div class="cta-row"><a href="/request-a-quote/" class="pill-btn pill-blue">Contact PanPanTech<span class="mono">→</span></a><a href="/products/" class="pill-btn pill-outline">Compare models</a></div></div>
</div></section>
<section class="band"><div class="container band-inner"><div class="gallery-intro"><div><p class="eyebrow">Product views</p><h2>{esc(product['model'])} photos and configuration views</h2></div><p>Selected product images from the latest PanPanTech specification pack and matched supplier materials.</p></div><div class="gallery">{gallery}</div></div></section>
<section class="band alt"><div class="container band-inner two-col"><div><p class="eyebrow">What it solves</p><h2 style="margin:0;font-size:38px;line-height:1.1;">Built around real site constraints</h2><div class="advantage-list">{advantages}</div></div>
<div class="panel"><p class="eyebrow">Best-fit applications</p><p style="margin:14px 0 0;font-size:16px;line-height:1.7;color:#3d4453;">Use these applications as the first filter, then send your route, floor material and operating schedule for model confirmation.</p><div class="chips">{applications}</div></div></div></section>
<section class="band"><div class="container band-inner"><p class="eyebrow">{esc(product['model'])}</p><h2 style="margin:0 0 42px;font-size:38px;">Technical specifications</h2><div class="spec-grid">{specs}</div><p style="margin:28px 0 0;font-family:'IBM Plex Mono',monospace;font-size:12px;color:#9aa3b5;">Specifications are based on the latest PanPanTech product specification document. Final configuration should be confirmed through the website contact form.</p></div></section>
<section class="dark-cta"><div class="container"><div><h2>Share your floor plan or material-flow route.</h2><p>For more product details, accessory options and deployment planning, contact PanPanTech through the website form or listed contact information.</p></div><a href="/request-a-quote/" class="pill-btn pill-blue">Open contact form<span class="mono">→</span></a></div></section>""",
        footer(),
    ]
    return "\n".join(body)


def legacy_pt90_page() -> str:
    return """<!doctype html><html><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1">
<title>ACR-0800 Autonomous Ride-On Floor Scrubber | PanPanTech</title>
<meta name="description" content="This legacy route now points to the PanPanTech ACR-0800 autonomous ride-on floor scrubber product page.">
<link rel="canonical" href="https://www.panpantechnology.com/products/acr-0800/">
<meta http-equiv="refresh" content="0; url=/products/acr-0800/">
<meta property="og:type" content="website"><meta property="og:title" content="ACR-0800 Autonomous Ride-On Floor Scrubber | PanPanTech"><meta property="og:description" content="The verified PanPanTech model name for this large-area scrubber page is ACR-0800."><meta property="og:url" content="https://www.panpantechnology.com/products/acr-0800/"><meta property="og:image" content="https://www.panpantechnology.com/assets/images/panpantech-social-card.jpg"><meta name="twitter:image" content="https://www.panpantechnology.com/assets/images/panpantech-social-card.jpg">
<script type="application/ld+json">""" + org_json_ld() + """</script></head><body><h1>ACR-0800 Autonomous Ride-On Floor Scrubber</h1><p><a href="/products/acr-0800/">Continue to ACR-0800 Autonomous Ride-On Floor Scrubber</a></p></body></html>"""


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8", newline="\n")


def update_sitemap(products: list[dict]) -> None:
    sitemap = ROOT / "sitemap.xml"
    content = sitemap.read_text(encoding="utf-8")
    product_urls = "\n".join(
        f"  <url><loc>{DOMAIN}/products/{product['slug']}/</loc><lastmod>2026-08-17</lastmod></url>"
        for product in products
    )
    content = re.sub(r"  <url><loc>https://www\.panpantechnology\.com/products/pt90/</loc>.*?</url>\n?", "", content, flags=re.S)
    content = re.sub(
        r"  <url><loc>https://www\.panpantechnology\.com/products/(?:acr-[^/]+|amr-[^/]+|asr-[^/]+|facade-cleaning-robot)/</loc><lastmod>2026-08-17</lastmod></url>\n?",
        "",
        content,
    )
    content = content.replace("</urlset>", product_urls + "\n</urlset>")
    sitemap.write_text(content, encoding="utf-8", newline="\n")


def main() -> None:
    products = extract_products()
    if not products:
        raise SystemExit("No products extracted from specification documents.")
    product_by_slug = {product["slug"]: product for product in products}
    missing = [name for category in CATEGORIES for name in category["products"] if name not in product_by_slug]
    if missing:
        raise SystemExit(f"Missing expected products: {missing}")

    write_text(ROOT / "products" / "index.html", products_index(products))
    for product in products:
        write_text(ROOT / "products" / product["slug"] / "index.html", product_page(product))
    write_text(ROOT / "products" / "pt90" / "index.html", legacy_pt90_page())
    update_sitemap(products)

    products_text = "\n".join((ROOT / "products" / product["slug"] / "index.html").read_text(encoding="utf-8") for product in products)
    for bad in BAD_MODEL_NAMES:
        if bad in products_text:
            raise SystemExit(f"Legacy supplier model name still present in generated product pages: {bad}")
    if re.search(r"\b(price|pricing|cost|priceCurrency|offers)\b", products_text, re.I):
        raise SystemExit("Generated product pages still contain price-related keywords.")


if __name__ == "__main__":
    main()
